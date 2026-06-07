"""Shared test helpers and fixtures for the readoc / readir suites.

The two CLIs live in ``bin/`` as extension-less scripts with a PEP 723 inline
dependency header. For *unit* tests we load them as importable modules with a
``SourceFileLoader`` (their heavy doc-parsing imports are lazy, so importing the
module is cheap). For *integration* tests we invoke them as subprocesses under
the current interpreter — the PEP 723 header is a plain comment to CPython, and
the test environment already has python-docx/openpyxl/pymupdf installed, so no
nested ``uv`` is needed.

Fixtures are generated on the fly into pytest's ``tmp_path``; nothing binary is
committed to the repo.
"""

import importlib.util
import subprocess
import sys
from importlib.machinery import SourceFileLoader
from pathlib import Path

import pytest

REPO_ROOT = Path(__file__).resolve().parent.parent
BIN = REPO_ROOT / "bin"


# --- Loading the scripts as modules (for unit tests) ---


def load_script(name):
    """Load bin/<name> as a module despite its lack of a .py extension."""
    path = BIN / name
    loader = SourceFileLoader(name, str(path))
    spec = importlib.util.spec_from_loader(name, loader)
    module = importlib.util.module_from_spec(spec)
    loader.exec_module(module)
    return module


@pytest.fixture(scope="session")
def readoc_mod():
    return load_script("readoc")


@pytest.fixture(scope="session")
def readir_mod():
    return load_script("readir")


@pytest.fixture(scope="session")
def readlib_mod():
    # The shared module carries a .py extension (it is imported, not run), so it
    # needs a slightly different load than the extension-less CLI scripts.
    path = BIN / "_readlib.py"
    loader = SourceFileLoader("_readlib", str(path))
    spec = importlib.util.spec_from_loader("_readlib", loader)
    module = importlib.util.module_from_spec(spec)
    loader.exec_module(module)
    return module


# --- Driving the CLIs end-to-end (for integration tests) ---


def run_cli(script, *args):
    """Run bin/<script> with the current interpreter; return (out, err, code)."""
    proc = subprocess.run(
        [sys.executable, str(BIN / script), *[str(a) for a in args]],
        capture_output=True,
        text=True,
    )
    return proc.stdout, proc.stderr, proc.returncode


@pytest.fixture
def readoc():
    def _run(*args):
        return run_cli("readoc", *args)

    return _run


@pytest.fixture
def readir():
    def _run(*args):
        return run_cli("readir", *args)

    return _run


# --- Fixture file builders ---

LONG_CELL = (
    "this is a deliberately long spreadsheet cell well beyond forty characters wide"
)


def make_docx(
    path,
    heading1="Main Title",
    heading2="A Subsection",
    body="Just a plain paragraph of body text.",
    table=True,
    comments=None,
):
    """Write a .docx with headings, a body paragraph, and (optionally) a table.

    ``comments`` is an optional list of ``(text, author)`` tuples; each is attached
    as a Word comment anchored to the body paragraph.
    """
    from docx import Document

    doc = Document()
    doc.add_heading(heading1, level=1)
    body_para = doc.add_paragraph(body)
    doc.add_heading(heading2, level=2)
    if table:
        t = doc.add_table(rows=2, cols=2)
        t.rows[0].cells[0].text = "Name"
        t.rows[0].cells[1].text = "Value"
        t.rows[1].cells[0].text = "alpha"
        t.rows[1].cells[1].text = "42"
    for text, author in comments or []:
        doc.add_comment(runs=body_para.runs, text=text, author=author)
    doc.save(str(path))
    return path


def make_xlsx(
    path, long_cell=LONG_CELL, multi_sheet=True, empty_sheet=True, comments=None
):
    """Write an .xlsx with a long cell, optional extra sheet, optional empty sheet.

    ``comments`` is an optional list of ``(coord, text, author)`` tuples written to
    the active "Data" sheet.
    """
    from openpyxl import Workbook
    from openpyxl.comments import Comment

    wb = Workbook()
    ws = wb.active
    ws.title = "Data"
    ws.append(["Header A", "Header B"])
    ws.append(["short", long_cell])
    ws.append([1, 2])
    if multi_sheet:
        ws2 = wb.create_sheet("Second")
        ws2.append(["x", "y"])
        ws2.append([10, 20])
    if empty_sheet:
        wb.create_sheet("Blank")
    for coord, text, author in comments or []:
        ws[coord].comment = Comment(text, author)
    wb.save(str(path))
    return path


def make_pdf(path, pages=("Hello from page one.",), comments=None):
    """Write a .pdf with one text block per entry in ``pages``.

    ``comments`` is an optional list of ``(page_index, text, author)`` tuples; each
    is added as a text (sticky-note) annotation on the given page.
    """
    import fitz

    doc = fitz.open()
    pages_list = []
    for text in pages:
        page = doc.new_page()
        page.insert_text((72, 72), text)
        pages_list.append(page)
    for page_index, text, author in comments or []:
        annot = pages_list[page_index].add_text_annot((200, 200), text)
        annot.set_info(title=author)
        annot.update()
    doc.save(str(path))
    doc.close()
    return path


# Distinctive paragraphs for context-window search tests. Only P2 mentions the
# search term "budget"; FARWORD lives two paragraphs away, so a ±1 paragraph
# window must exclude it while still including its immediate neighbours.
PROSE_PARAS = [
    "Intro paragraph about onboarding new staff.",
    "The budget paragraph discusses overruns in detail.",
    "Hiring plans are covered in this neighbour paragraph.",
    "FARWORD marks a distant paragraph that should stay hidden.",
]


def make_prose_docx(path, paras=PROSE_PARAS, heading="Quarterly Report"):
    """Write a multi-paragraph .docx for paragraph/char context search tests."""
    from docx import Document

    doc = Document()
    doc.add_heading(heading, level=1)
    for p in paras:
        doc.add_paragraph(p)
    doc.save(str(path))
    return path


def make_search_xlsx(path):
    """Write an .xlsx with known coordinates: headers in row 1, row headers in
    column A, the term 'overrun' at C3 (Budget!C3), and a cell comment on B3."""
    from openpyxl import Workbook
    from openpyxl.comments import Comment

    wb = Workbook()
    ws = wb.active
    ws.title = "Budget"
    ws.append(["Quarter", "Amount", "Notes"])
    ws.append(["Q1", 1000, "on track"])
    ws.append(["Q2", 2500, "budget overrun here"])
    ws.append(["Q3", 1800, "stable"])
    ws["B3"].comment = Comment("double-check this amount", "Reviewer")
    wb.save(str(path))
    return path


@pytest.fixture
def prose_docx(tmp_path):
    return make_prose_docx(tmp_path / "report.docx")


@pytest.fixture
def search_xlsx(tmp_path):
    return make_search_xlsx(tmp_path / "sheet.xlsx")


@pytest.fixture
def search_tree(tmp_path):
    """A folder mixing a multi-paragraph docx, a coordinate-known xlsx, and a
    single-long-line txt — the substrate for the readir search-mode tests."""
    root = tmp_path / "searchtree"
    root.mkdir()
    make_prose_docx(root / "report.docx")
    make_search_xlsx(root / "sheet.xlsx")
    # One long physical line so --context-chars has room to window around a match.
    (root / "blob.txt").write_text(
        "alpha beta gamma delta the budget line continues onward epsilon zeta eta " * 3,
        encoding="utf-8",
    )
    return root


@pytest.fixture
def docx_file(tmp_path):
    return make_docx(tmp_path / "doc.docx")


@pytest.fixture
def xlsx_file(tmp_path):
    return make_xlsx(tmp_path / "sheet.xlsx")


@pytest.fixture
def pdf_file(tmp_path):
    return make_pdf(tmp_path / "single.pdf")


@pytest.fixture
def docx_with_comments(tmp_path):
    return make_docx(
        tmp_path / "commented.docx",
        comments=[("First reviewer note", "Alice"), ("Second note", "Bob")],
    )


@pytest.fixture
def xlsx_with_comments(tmp_path):
    return make_xlsx(
        tmp_path / "commented.xlsx",
        comments=[("A1", "Header looks off", "Carol")],
    )


@pytest.fixture
def pdf_with_comments(tmp_path):
    return make_pdf(
        tmp_path / "commented.pdf",
        comments=[(0, "A sticky note here", "Dave")],
    )


@pytest.fixture
def multipage_pdf(tmp_path):
    return make_pdf(
        tmp_path / "multi.pdf",
        pages=("Content of the first page.", "Content of the second page."),
    )


@pytest.fixture
def sample_tree(tmp_path):
    """A nested folder mixing readable text, Office docs, an unsupported file,
    and one oversized text file — the substrate for the readir integration tests.

    Returns the root path. Notable contents:
      - top.md / data.csv / config.json   (readable text)
      - notes.txt with the word "budget"  (search target)
      - sub/report.docx mentioning "BUDGET" in a heading (case-insensitive search)
      - sub/big.txt  ~3 KB                 (oversized for --max-size 1)
      - sub/deeper/deepfile.txt            (depth-2 file, for --max-depth tests)
      - image.png                          (unsupported format)
    """
    root = tmp_path / "tree"
    deeper = root / "sub" / "deeper"
    deeper.mkdir(parents=True)
    sub = root / "sub"

    (root / "top.md").write_text(
        "# Top\n\nSome markdown about budgets.\n", encoding="utf-8"
    )
    (root / "data.csv").write_text("a,b,c\n1,2,3\n", encoding="utf-8")
    (root / "config.json").write_text('{"key": "value"}\n', encoding="utf-8")
    (root / "notes.txt").write_text(
        "Line one\nThe budget is tight\nLine three\n", encoding="utf-8"
    )
    (root / "image.png").write_bytes(b"\x89PNG\r\n\x1a\n" + b"\x00" * 32)

    make_docx(
        sub / "report.docx",
        heading1="BUDGET Overview",
        body="Detailed budget discussion follows.",
        table=False,
    )
    (sub / "big.txt").write_text("padding line\n" * 300, encoding="utf-8")  # ~3.6 KB

    (deeper / "deepfile.txt").write_text("nothing special here\n", encoding="utf-8")

    return root
