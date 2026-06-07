"""Tests for the editdoc CLI (Edit-tool-equivalent for .docx / .xlsx).

editdoc reads a JSON edit spec (a single object or an array of objects) from
stdin and applies it to the target file given as an argv. Fixtures are built on
the fly with python-docx / openpyxl; results are verified by re-opening the file.
"""

import hashlib

import docx
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import Workbook, load_workbook


# --- Fixture builders specific to editing tests ---


def add_hyperlink_run(paragraph, text, url="https://example.com"):
    """Append a real <w:hyperlink> (with a nested run) to ``paragraph``.

    python-docx's ``paragraph.runs`` does NOT include runs nested inside a
    hyperlink, so this builds the exact shape that a naive run-only rewrite would
    silently scramble — the substrate for the hyperlink-safety tests.
    """
    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def make_runs_docx(path, runs, style=None):
    """Write a .docx whose single body paragraph is composed of explicit runs.

    ``runs`` is a list of ``(text, bold)`` tuples, so tests can build a paragraph
    that is deliberately fragmented across runs with mixed formatting — the case
    that defeats a naive find/replace on the raw XML.
    """
    doc = Document()
    p = doc.add_paragraph()
    if style:
        p.style = style
    for text, bold in runs:
        r = p.add_run(text)
        r.bold = bold
    doc.save(str(path))
    return path


def make_paras_docx(path, paras):
    """Write a .docx with one plain paragraph per string in ``paras``."""
    doc = Document()
    for text in paras:
        doc.add_paragraph(text)
    doc.save(str(path))
    return path


def make_xlsx(path, rows=(("Quarter", "Amount"), ("Q1", 1000)), sheet="Budget"):
    """Write a small .xlsx with the given rows on a named sheet."""
    wb = Workbook()
    ws = wb.active
    ws.title = sheet
    for row in rows:
        ws.append(list(row))
    wb.save(str(path))
    return path


def make_xlsx_with_chart(path, sheet="Budget"):
    """Write an .xlsx that embeds a bar chart — a feature openpyxl drops on save,
    so editdoc must refuse to edit it without --force."""
    from openpyxl.chart import BarChart, Reference

    wb = Workbook()
    ws = wb.active
    ws.title = sheet
    ws.append(["Quarter", "Amount"])
    ws.append(["Q1", 1000])
    ws.append(["Q2", 1500])
    chart = BarChart()
    chart.add_data(
        Reference(ws, min_col=2, min_row=1, max_row=3), titles_from_data=True
    )
    ws.add_chart(chart, "E5")
    wb.save(str(path))
    return path


def cell_value(path, sheet, coord):
    wb = load_workbook(str(path))
    return wb[sheet][coord].value


def body_texts(path):
    """Return the visible text of every body paragraph in document order."""
    return [p.text for p in Document(str(path)).paragraphs]


def sha256(path):
    return hashlib.sha256(path.read_bytes()).hexdigest()


# --- Task 1: skeleton (argv, stdin JSON, dispatch, error paths) ---


def test_no_args_shows_usage():
    # editdoc with zero argv should print usage and exit 1; reuse run_cli directly.
    from conftest import run_cli

    out, err, code = run_cli("editdoc", stdin="{}")
    assert code == 1
    assert "Usage" in out or "Usage" in err


def test_missing_file_errors(editdoc, tmp_path):
    out, err, code = editdoc(
        tmp_path / "nope.docx", {"old_string": "a", "new_string": "b"}
    )
    assert code == 1
    assert "not found" in err.lower()


def test_unsupported_extension_errors(editdoc, tmp_path):
    f = tmp_path / "x.pdf"
    f.write_text("hi")
    out, err, code = editdoc(f, {"old_string": "a", "new_string": "b"})
    assert code == 1
    assert "unsupported" in err.lower()


def test_invalid_json_errors(editdoc, tmp_path):
    f = make_paras_docx(tmp_path / "x.docx", ["hello world"])
    out, err, code = editdoc(f, "{bad json")
    assert code == 1
    assert "json" in err.lower()


def test_empty_spec_is_noop_success(editdoc, tmp_path):
    # An empty array of edits is a valid no-op: nothing to do, exit 0, file intact.
    f = make_paras_docx(tmp_path / "x.docx", ["hello world"])
    before = sha256(f)
    out, err, code = editdoc(f, [])
    assert code == 0, err
    assert sha256(f) == before


# --- Task 2: docx in-paragraph exact-match replace ---


def test_docx_basic_replace(editdoc, tmp_path):
    f = make_paras_docx(tmp_path / "x.docx", ["The report is due Monday."])
    out, err, code = editdoc(f, {"old_string": "Monday", "new_string": "Friday"})
    assert code == 0, err
    assert body_texts(f) == ["The report is due Friday."]
    assert "✓" in out  # success line


def test_docx_replace_is_whitespace_exact(editdoc, tmp_path):
    # A near-match with different internal spacing must NOT match.
    f = make_paras_docx(tmp_path / "x.docx", ["The  report  is due."])  # two spaces
    out, err, code = editdoc(
        f,
        {"old_string": "The report is", "new_string": "X"},  # single spaces
    )
    assert code == 1
    assert "not found" in err.lower()
    assert body_texts(f) == ["The  report  is due."]  # untouched


def test_docx_replace_across_fragmented_runs_inherits_first_run_format(
    editdoc, tmp_path
):
    # Visible text "The report is due" is split across three runs, the middle one
    # bold. Replacing a span that starts in the first (non-bold) run must succeed
    # and the replacement must inherit the FIRST spanned run's formatting.
    f = make_runs_docx(
        tmp_path / "x.docx",
        runs=[("The ", False), ("report", True), (" is due", False)],
    )
    out, err, code = editdoc(
        f, {"old_string": "The report is due", "new_string": "The memo is ready"}
    )
    assert code == 0, err
    doc = Document(str(f))
    para = doc.paragraphs[0]
    assert para.text == "The memo is ready"
    # The run carrying the replacement text inherits the first run's format (not bold).
    carrier = next(r for r in para.runs if "memo" in r.text)
    assert carrier.bold in (False, None)


def test_docx_replace_preserves_untouched_run_formatting(editdoc, tmp_path):
    # Replacing text wholly inside the first run leaves the bold middle run intact.
    f = make_runs_docx(
        tmp_path / "x.docx",
        runs=[("The ", False), ("report", True), (" is due", False)],
    )
    out, err, code = editdoc(f, {"old_string": "The ", "new_string": "A "})
    assert code == 0, err
    doc = Document(str(f))
    para = doc.paragraphs[0]
    assert para.text == "A report is due"
    bold_run = next(r for r in para.runs if r.text == "report")
    assert bold_run.bold is True


def test_docx_replace_not_found_errors(editdoc, tmp_path):
    f = make_paras_docx(tmp_path / "x.docx", ["hello world"])
    out, err, code = editdoc(f, {"old_string": "goodbye", "new_string": "hi"})
    assert code == 1
    assert "not found" in err.lower()
    assert body_texts(f) == ["hello world"]


def test_docx_replace_ambiguous_without_replace_all_errors(editdoc, tmp_path):
    f = make_paras_docx(
        tmp_path / "x.docx", ["alpha here", "alpha there"]
    )  # 'alpha' twice
    out, err, code = editdoc(f, {"old_string": "alpha", "new_string": "beta"})
    assert code == 1
    assert "2" in err  # mentions the count
    assert "replace_all" in err or "context" in err.lower()
    assert body_texts(f) == ["alpha here", "alpha there"]  # untouched


def test_docx_replace_all(editdoc, tmp_path):
    f = make_paras_docx(tmp_path / "x.docx", ["alpha here", "alpha there"])
    out, err, code = editdoc(
        f, {"old_string": "alpha", "new_string": "beta", "replace_all": True}
    )
    assert code == 0, err
    assert body_texts(f) == ["beta here", "beta there"]


def test_docx_replace_disambiguate_with_context(editdoc, tmp_path):
    # Adding surrounding context makes an otherwise-duplicated word unique.
    f = make_paras_docx(tmp_path / "x.docx", ["alpha here", "alpha there"])
    out, err, code = editdoc(f, {"old_string": "alpha here", "new_string": "beta here"})
    assert code == 0, err
    assert body_texts(f) == ["beta here", "alpha there"]


def test_docx_old_string_with_newline_errors(editdoc, tmp_path):
    f = make_paras_docx(tmp_path / "x.docx", ["line one", "line two"])
    out, err, code = editdoc(f, {"old_string": "line one\nline two", "new_string": "x"})
    assert code == 1
    assert "newline" in err.lower() or "paragraph" in err.lower()


def test_docx_replace_inside_table_cell(editdoc, tmp_path):
    # The edit surface includes table cells.
    doc = Document()
    t = doc.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "Name"
    t.rows[0].cells[1].text = "old value"
    f = tmp_path / "x.docx"
    doc.save(str(f))
    out, err, code = editdoc(f, {"old_string": "old value", "new_string": "new value"})
    assert code == 0, err
    cell_text = Document(str(f)).tables[0].rows[0].cells[1].text
    assert cell_text == "new value"


# --- Task 3: docx paragraph / range replace ---


def test_docx_single_paragraph_replace(editdoc, tmp_path):
    f = make_paras_docx(
        tmp_path / "x.docx", ["Keep me", "Quarterly results were strong", "Keep me too"]
    )
    out, err, code = editdoc(
        f, {"start_contains": "Quarterly results", "new_text": "Replaced line"}
    )
    assert code == 0, err
    assert body_texts(f) == ["Keep me", "Replaced line", "Keep me too"]


def test_docx_paragraph_split_into_many(editdoc, tmp_path):
    # new_text with newlines turns one paragraph into several.
    f = make_paras_docx(tmp_path / "x.docx", ["before", "TARGET para", "after"])
    out, err, code = editdoc(
        f, {"start_contains": "TARGET", "new_text": "one\ntwo\nthree"}
    )
    assert code == 0, err
    assert body_texts(f) == ["before", "one", "two", "three", "after"]


def test_docx_range_replace_five_to_three(editdoc, tmp_path):
    f = make_paras_docx(
        tmp_path / "x.docx",
        ["head", "P1 start", "P2", "P3", "P4", "P5 end", "tail"],
    )
    out, err, code = editdoc(
        f,
        {
            "start_contains": "P1 start",
            "end_contains": "P5 end",
            "new_text": "new A\nnew B\nnew C",
        },
    )
    assert code == 0, err
    assert body_texts(f) == ["head", "new A", "new B", "new C", "tail"]


def test_docx_range_delete_with_empty_new_text(editdoc, tmp_path):
    f = make_paras_docx(
        tmp_path / "x.docx", ["keep", "drop 1", "drop 2", "drop 3", "keep too"]
    )
    out, err, code = editdoc(
        f,
        {"start_contains": "drop 1", "end_contains": "drop 3", "new_text": ""},
    )
    assert code == 0, err
    assert body_texts(f) == ["keep", "keep too"]


def test_docx_single_paragraph_delete(editdoc, tmp_path):
    f = make_paras_docx(tmp_path / "x.docx", ["keep", "remove me", "keep too"])
    out, err, code = editdoc(f, {"start_contains": "remove me", "new_text": ""})
    assert code == 0, err
    assert body_texts(f) == ["keep", "keep too"]


def test_docx_paragraph_split_inherits_style(editdoc, tmp_path):
    # Splitting a styled paragraph: every produced paragraph keeps the style.
    doc = Document()
    doc.add_paragraph("plain")
    doc.add_paragraph("Heading target", style="Heading 1")
    f = tmp_path / "x.docx"
    doc.save(str(f))
    out, err, code = editdoc(
        f, {"start_contains": "Heading target", "new_text": "H one\nH two"}
    )
    assert code == 0, err
    styled = [
        p.text for p in Document(str(f)).paragraphs if p.style.name == "Heading 1"
    ]
    assert styled == ["H one", "H two"]


def test_docx_paragraph_anchor_not_unique_errors(editdoc, tmp_path):
    f = make_paras_docx(tmp_path / "x.docx", ["dup line", "dup line"])
    out, err, code = editdoc(f, {"start_contains": "dup line", "new_text": "x"})
    assert code == 1
    assert "2" in err
    assert body_texts(f) == ["dup line", "dup line"]


def test_docx_paragraph_anchor_not_found_errors(editdoc, tmp_path):
    f = make_paras_docx(tmp_path / "x.docx", ["a", "b"])
    out, err, code = editdoc(f, {"start_contains": "nope", "new_text": "x"})
    assert code == 1
    assert "not found" in err.lower()


def test_docx_range_end_before_start_errors(editdoc, tmp_path):
    f = make_paras_docx(tmp_path / "x.docx", ["first", "second"])
    out, err, code = editdoc(
        f,
        {"start_contains": "second", "end_contains": "first", "new_text": "x"},
    )
    assert code == 1
    assert "before" in err.lower() or "order" in err.lower()


def test_docx_range_cross_container_errors(editdoc, tmp_path):
    doc = Document()
    doc.add_paragraph("start here in body")
    t = doc.add_table(rows=1, cols=1)
    t.rows[0].cells[0].text = "end here in cell"
    f = tmp_path / "x.docx"
    doc.save(str(f))
    out, err, code = editdoc(
        f,
        {
            "start_contains": "start here",
            "end_contains": "end here",
            "new_text": "x",
        },
    )
    assert code == 1
    assert "container" in err.lower() or "same" in err.lower()


# --- Task 4: xlsx cell edit ---


def test_xlsx_set_existing_cell(editdoc, tmp_path):
    f = make_xlsx(tmp_path / "x.xlsx", rows=(("Quarter", "Amount"), ("Q1", 1000)))
    out, err, code = editdoc(f, {"sheet": "Budget", "cell": "B2", "value": "1250"})
    assert code == 0, err
    assert cell_value(f, "Budget", "B2") == 1250  # numeric coercion
    assert "Budget!B2" in out


def test_xlsx_set_empty_cell(editdoc, tmp_path):
    f = make_xlsx(tmp_path / "x.xlsx", rows=(("Quarter", "Amount"), ("Q1", 1000)))
    out, err, code = editdoc(f, {"sheet": "Budget", "cell": "C5", "value": "hello"})
    assert code == 0, err
    assert cell_value(f, "Budget", "C5") == "hello"


def test_xlsx_numeric_vs_string_coercion(editdoc, tmp_path):
    f = make_xlsx(tmp_path / "x.xlsx")
    editdoc(f, {"sheet": "Budget", "cell": "A5", "value": "3.5"})
    editdoc(f, {"sheet": "Budget", "cell": "A6", "value": "N/A"})
    assert cell_value(f, "Budget", "A5") == 3.5
    assert cell_value(f, "Budget", "A6") == "N/A"


def test_xlsx_numeric_passed_as_json_number(editdoc, tmp_path):
    f = make_xlsx(tmp_path / "x.xlsx")
    out, err, code = editdoc(f, {"sheet": "Budget", "cell": "A7", "value": 42})
    assert code == 0, err
    assert cell_value(f, "Budget", "A7") == 42


def test_xlsx_bad_sheet_errors(editdoc, tmp_path):
    f = make_xlsx(tmp_path / "x.xlsx")
    out, err, code = editdoc(f, {"sheet": "Nonexistent", "cell": "A1", "value": "x"})
    assert code == 1
    assert "sheet" in err.lower()


def test_xlsx_bad_coordinate_errors(editdoc, tmp_path):
    f = make_xlsx(tmp_path / "x.xlsx")
    out, err, code = editdoc(f, {"sheet": "Budget", "cell": "1", "value": "x"})
    assert code == 1
    assert "cell" in err.lower() or "coordinate" in err.lower()


def test_xlsx_missing_value_errors(editdoc, tmp_path):
    f = make_xlsx(tmp_path / "x.xlsx")
    out, err, code = editdoc(f, {"sheet": "Budget", "cell": "A1"})
    assert code == 1


# --- Task 5: batch atomicity (validate-all-then-apply) ---


def test_docx_batch_all_or_nothing_on_failure(editdoc, tmp_path):
    # A valid edit followed by an invalid one must leave the file byte-identical:
    # the first edit's (in-memory) change is never written when a later edit fails.
    f = make_paras_docx(tmp_path / "x.docx", ["alpha unique", "beta line"])
    before = sha256(f)
    spec = [
        {"old_string": "alpha unique", "new_string": "ALPHA"},  # valid
        {"old_string": "does-not-exist", "new_string": "x"},  # invalid
    ]
    out, err, code = editdoc(f, spec)
    assert code == 1
    assert "edit 2" in err
    assert sha256(f) == before  # nothing written


def test_docx_batch_all_valid_applies_every_edit(editdoc, tmp_path):
    f = make_paras_docx(tmp_path / "x.docx", ["alpha", "beta", "gamma"])
    spec = [
        {"old_string": "alpha", "new_string": "A"},
        {"old_string": "gamma", "new_string": "G"},
    ]
    out, err, code = editdoc(f, spec)
    assert code == 0, err
    assert body_texts(f) == ["A", "beta", "G"]


def test_docx_batch_multiple_range_deletes(editdoc, tmp_path):
    # Two independent range deletes in one batch both succeed; deleting an earlier
    # block must not invalidate the second block's anchors.
    f = make_paras_docx(
        tmp_path / "x.docx",
        ["keep0", "d1a", "d1b", "keep1", "d2a", "d2b", "keep2"],
    )
    spec = [
        {"start_contains": "d1a", "end_contains": "d1b", "new_text": ""},
        {"start_contains": "d2a", "end_contains": "d2b", "new_text": ""},
    ]
    out, err, code = editdoc(f, spec)
    assert code == 0, err
    assert body_texts(f) == ["keep0", "keep1", "keep2"]


def test_xlsx_batch_all_or_nothing_on_failure(editdoc, tmp_path):
    f = make_xlsx(tmp_path / "x.xlsx")
    before = sha256(f)
    spec = [
        {"sheet": "Budget", "cell": "A5", "value": "1"},  # valid
        {"sheet": "Nonexistent", "cell": "A1", "value": "x"},  # invalid
    ]
    out, err, code = editdoc(f, spec)
    assert code == 1
    assert sha256(f) == before


# --- Review fixes: hyperlink safety, merged cells, empty anchors ---


def test_docx_edit_paragraph_with_hyperlink_refuses_loudly(editdoc, tmp_path):
    # A paragraph whose text includes a hyperlink (whose runs are invisible to
    # paragraph.runs) must NOT be rewritten — that silently scrambles it. Refuse
    # loudly and leave the file byte-identical.
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("See the ")
    add_hyperlink_run(p, "docs")
    p.add_run(" for details on Monday.")
    f = tmp_path / "x.docx"
    doc.save(str(f))
    before = sha256(f)
    out, err, code = editdoc(f, {"old_string": "Monday", "new_string": "Friday"})
    assert code == 1
    assert "hyperlink" in err.lower() or "safely" in err.lower()
    assert sha256(f) == before  # nothing scrambled


def test_docx_hyperlink_in_other_paragraph_does_not_block_edit(editdoc, tmp_path):
    # The guard is scoped to the edited paragraph: a hyperlink elsewhere in the
    # document must not block an edit to a clean paragraph.
    doc = Document()
    hp = doc.add_paragraph()
    hp.add_run("Visit ")
    add_hyperlink_run(hp, "our site")
    doc.add_paragraph("Plain paragraph with target word.")
    f = tmp_path / "x.docx"
    doc.save(str(f))
    out, err, code = editdoc(f, {"old_string": "target word", "new_string": "FIXED"})
    assert code == 0, err
    texts = [p.text for p in Document(str(f)).paragraphs]
    assert "Plain paragraph with FIXED." in texts
    # The hyperlinked paragraph is untouched and intact.
    assert any("Visit our site" == t for t in texts)


def test_docx_paragraph_op_on_hyperlink_paragraph_refuses(editdoc, tmp_path):
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Anchor ")
    add_hyperlink_run(p, "linked")
    p.add_run(" tail")
    f = tmp_path / "x.docx"
    doc.save(str(f))
    before = sha256(f)
    out, err, code = editdoc(f, {"start_contains": "Anchor", "new_text": "replaced"})
    assert code == 1
    assert "hyperlink" in err.lower() or "safely" in err.lower()
    assert sha256(f) == before


def test_docx_replace_in_merged_cell(editdoc, tmp_path):
    # A horizontally-merged cell is one logical cell; a unique string in it must
    # not be reported as matching twice.
    doc = Document()
    t = doc.add_table(rows=1, cols=2)
    merged = t.rows[0].cells[0].merge(t.rows[0].cells[1])
    merged.text = "MERGED unique target"
    f = tmp_path / "x.docx"
    doc.save(str(f))
    out, err, code = editdoc(
        f, {"old_string": "MERGED unique target", "new_string": "done"}
    )
    assert code == 0, err
    assert Document(str(f)).tables[0].rows[0].cells[0].text == "done"


def test_docx_empty_start_contains_rejected(editdoc, tmp_path):
    f = make_paras_docx(tmp_path / "x.docx", ["only para"])
    before = sha256(f)
    out, err, code = editdoc(f, {"start_contains": "", "new_text": "HIJACKED"})
    assert code == 1
    assert body_texts(f) == ["only para"]
    assert sha256(f) == before


def test_docx_empty_end_contains_rejected(editdoc, tmp_path):
    f = make_paras_docx(tmp_path / "x.docx", ["alpha", "beta"])
    out, err, code = editdoc(
        f, {"start_contains": "alpha", "end_contains": "", "new_text": "x"}
    )
    assert code == 1


# --- Review round 2: xlsx coercion edges + lossy-feature detection ---


def test_xlsx_leading_zero_kept_as_string(editdoc, tmp_path):
    # Zip/account codes must not be truncated to ints.
    f = make_xlsx(tmp_path / "x.xlsx")
    out, err, code = editdoc(f, {"sheet": "Budget", "cell": "A5", "value": "00501"})
    assert code == 0, err
    assert cell_value(f, "Budget", "A5") == "00501"


def test_xlsx_nan_inf_kept_as_string_not_blanked(editdoc, tmp_path):
    # float("nan")/("inf") parse but openpyxl can't store them — they must NOT
    # silently blank the cell. Keep them as text.
    f = make_xlsx(tmp_path / "x.xlsx")
    for cell, val in (("A5", "nan"), ("A6", "inf"), ("A7", "-inf")):
        out, err, code = editdoc(f, {"sheet": "Budget", "cell": cell, "value": val})
        assert code == 0, err
    assert cell_value(f, "Budget", "A5") == "nan"
    assert cell_value(f, "Budget", "A6") == "inf"
    assert cell_value(f, "Budget", "A7") == "-inf"


def test_xlsx_scientific_and_underscore_kept_as_string(editdoc, tmp_path):
    f = make_xlsx(tmp_path / "x.xlsx")
    editdoc(f, {"sheet": "Budget", "cell": "A5", "value": "1e3"})
    editdoc(f, {"sheet": "Budget", "cell": "A6", "value": "1_000"})
    assert cell_value(f, "Budget", "A5") == "1e3"
    assert cell_value(f, "Budget", "A6") == "1_000"


def test_xlsx_plain_integer_still_coerced(editdoc, tmp_path):
    # The canonical-round-trip rule must still coerce ordinary numbers.
    f = make_xlsx(tmp_path / "x.xlsx")
    editdoc(f, {"sheet": "Budget", "cell": "A5", "value": "1250"})
    editdoc(f, {"sheet": "Budget", "cell": "A6", "value": "-42"})
    editdoc(f, {"sheet": "Budget", "cell": "A7", "value": "3.5"})
    assert cell_value(f, "Budget", "A5") == 1250
    assert cell_value(f, "Budget", "A6") == -42
    assert cell_value(f, "Budget", "A7") == 3.5


def test_xlsx_null_clears_cell(editdoc, tmp_path):
    f = make_xlsx(tmp_path / "x.xlsx", rows=(("Quarter", "Amount"), ("Q1", 1000)))
    out, err, code = editdoc(f, {"sheet": "Budget", "cell": "B2", "value": None})
    assert code == 0, err
    assert cell_value(f, "Budget", "B2") is None


def test_xlsx_bool_stored(editdoc, tmp_path):
    f = make_xlsx(tmp_path / "x.xlsx")
    out, err, code = editdoc(f, {"sheet": "Budget", "cell": "A5", "value": True})
    assert code == 0, err
    assert cell_value(f, "Budget", "A5") is True


def test_xlsx_chart_workbook_refused_without_force(editdoc, tmp_path):
    f = make_xlsx_with_chart(tmp_path / "x.xlsx")
    before = sha256(f)
    out, err, code = editdoc(f, {"sheet": "Budget", "cell": "A2", "value": "Q1x"})
    assert code == 1
    assert "chart" in err.lower()
    assert "--force" in err or "force" in err.lower()
    assert sha256(f) == before  # not rewritten, chart preserved


def test_xlsx_chart_workbook_editable_with_force(editdoc, tmp_path):
    f = make_xlsx_with_chart(tmp_path / "x.xlsx")
    out, err, code = editdoc(
        f, {"sheet": "Budget", "cell": "A2", "value": "Q1x"}, "--force"
    )
    assert code == 0, err
    assert cell_value(f, "Budget", "A2") == "Q1x"


def test_xlsx_plain_workbook_not_refused(editdoc, tmp_path):
    # A chart-free data sheet edits with no friction and no flag.
    f = make_xlsx(tmp_path / "x.xlsx")
    out, err, code = editdoc(f, {"sheet": "Budget", "cell": "A5", "value": "ok"})
    assert code == 0, err


def test_unknown_flag_rejected(tmp_path):
    from conftest import run_cli

    f = make_paras_docx(tmp_path / "x.docx", ["hello"])
    out, err, code = run_cli(
        "editdoc", f, "--frce", stdin='{"old_string": "hello", "new_string": "hi"}'
    )
    assert code == 1
    assert "flag" in err.lower() or "unknown" in err.lower()


def test_multiple_files_rejected(tmp_path):
    # editdoc edits a single file; passing more than one must error, not silently
    # edit only the first.
    from conftest import run_cli

    a = make_paras_docx(tmp_path / "a.docx", ["one"])
    b = make_paras_docx(tmp_path / "b.docx", ["two"])
    before_b = sha256(b)
    out, err, code = run_cli(
        "editdoc",
        a,
        b,
        stdin='{"old_string": "one", "new_string": "X"}',
    )
    assert code == 1
    assert "one file" in err.lower() or "single" in err.lower()
    assert sha256(b) == before_b  # second file untouched
