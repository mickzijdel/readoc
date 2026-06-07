"""Shared library for the readoc / readir CLIs.

This is a plain importable module — NOT a runnable script — so it carries no PEP
723 header. The third-party libraries it touches (python-docx, openpyxl, pymupdf)
are imported lazily inside the functions that need them, and are provided by
whichever CLI imports this module (each declares them in its own PEP 723 header).

Both bin/readoc and bin/readir locate this file via
``os.path.dirname(os.path.realpath(__file__))`` — realpath resolves the PATH /
skills symlinks back to the real bin/ directory, so a bare `import _readlib`
works whether the script is run directly, through a symlink, or under uv.
"""

import io
import os
import re
import sys


# --- Sizes ---


def human_size(size_bytes):
    for unit in ("B", "KB", "MB", "GB"):
        if abs(size_bytes) < 1024:
            if unit == "B":
                return f"{size_bytes} {unit}"
            return f"{size_bytes:.1f} {unit}"
        size_bytes /= 1024
    return f"{size_bytes:.1f} TB"


# --- Comment extraction ---


def comments_block(comments):
    """Render a trailing Comments section, or '' when there are no comments."""
    if not comments:
        return ""
    return "\n--- Comments ---\n" + "\n".join(comments)


def docx_comments(doc):
    """Return formatted comment lines from a python-docx Document, or []."""
    try:
        return [f"{(c.author or 'Unknown')}: {c.text}".rstrip() for c in doc.comments]
    except Exception:
        return []


def xlsx_comments(path):
    """Return formatted cell-comment lines. Requires a non-read_only load, since
    read-only cells expose no comment attribute."""
    from openpyxl import load_workbook

    lines = []
    try:
        wb = load_workbook(path, read_only=False, data_only=True)
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            for row in ws.iter_rows():
                for cell in row:
                    comment = cell.comment
                    if comment and (comment.text or "").strip():
                        author = (comment.author or "Unknown").strip() or "Unknown"
                        text = comment.text.strip()
                        lines.append(
                            f"[{sheet_name}!{cell.coordinate}] {author}: {text}"
                        )
        wb.close()
    except Exception:
        return []
    return lines


def pdf_comments(doc):
    """Return formatted annotation-comment lines from a PyMuPDF document, or []."""
    lines = []
    try:
        for page_num, page in enumerate(doc, 1):
            for annot in page.annots() or []:
                info = annot.info
                content = (info.get("content") or "").strip()
                if not content:
                    continue
                author = (info.get("title") or "Unknown").strip() or "Unknown"
                lines.append(f"[Page {page_num}] {author}: {content}")
    except Exception:
        return []
    return lines


# --- Text-returning readers ---


def read_pdf_text(path, include_comments=True):
    import fitz

    lines = []
    doc = fitz.open(path)
    for page_num, page in enumerate(doc, 1):
        text = page.get_text().strip()
        if text:
            if len(doc) > 1:
                lines.append(f"\n--- Page {page_num} ---\n")
            lines.append(text)
    block = comments_block(pdf_comments(doc)) if include_comments else ""
    doc.close()
    text = "\n".join(lines)
    return text + ("\n" + block if block else "")


def read_docx_text(path, include_comments=True):
    from docx import Document

    lines = []
    doc = Document(path)
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            if para.style and para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading ", "").strip()
                try:
                    prefix = "#" * int(level) + " "
                except ValueError:
                    prefix = "# "
                lines.append(f"{prefix}{text}")
            else:
                lines.append(text)
    if doc.tables:
        for i, table in enumerate(doc.tables):
            lines.append(f"\n--- Table {i + 1} ---")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                lines.append(" | ".join(cells))
    if include_comments:
        block = comments_block(docx_comments(doc))
        if block:
            lines.append(block)
    return "\n".join(lines)


def read_xlsx_text(path, include_comments=True):
    from openpyxl import load_workbook

    buf = io.StringIO()
    wb = load_workbook(path, read_only=True, data_only=True)
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        buf.write(f"=== Sheet: {sheet_name} ===\n\n")
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            buf.write("(empty sheet)\n\n")
            continue
        str_rows = []
        for row in rows:
            str_rows.append([str(cell) if cell is not None else "" for cell in row])
        col_widths = []
        for col_idx in range(len(str_rows[0])):
            max_width = 0
            for row in str_rows:
                if col_idx < len(row):
                    max_width = max(max_width, len(row[col_idx]))
            # Cap the *padding* width for readability, but never truncate content.
            col_widths.append(min(max_width, 40))
        for row_idx, row in enumerate(str_rows):
            padded = []
            for col_idx, cell in enumerate(row):
                width = col_widths[col_idx] if col_idx < len(col_widths) else 10
                # ljust pads short cells for alignment; long cells overflow in full (no truncation).
                padded.append(cell.ljust(width))
            buf.write(" | ".join(padded) + "\n")
            if row_idx == 0:
                buf.write("-+-".join("-" * w for w in col_widths) + "\n")
        buf.write("\n")
    wb.close()
    result = buf.getvalue()
    if include_comments:
        block = comments_block(xlsx_comments(path))
        if block:
            result += block + "\n"
    return result


def read_text_file(path):
    for encoding in ("utf-8", "latin-1"):
        try:
            with open(path, "r", encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    return None


# --- Supported extensions ---

READABLE_TEXT = {".md", ".txt", ".csv", ".json", ".yaml", ".yml"}
READABLE_SPECIAL = {".docx", ".xlsx", ".pdf"}
READABLE_ALL = READABLE_TEXT | READABLE_SPECIAL


def read_file_content(path, include_comments=True):
    """Read a file and return its text content, or None on failure."""
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext in READABLE_TEXT:
            return read_text_file(path)
        elif ext == ".docx":
            return read_docx_text(path, include_comments)
        elif ext == ".xlsx":
            return read_xlsx_text(path, include_comments)
        elif ext == ".pdf":
            return read_pdf_text(path, include_comments)
    except Exception as e:
        print(f"  Warning: Failed to read {path}: {e}", file=sys.stderr)
    return None


# --- Search (structure-aware: lines, paragraphs, chars, and spreadsheet cells) ---

_PARA_SPLIT = re.compile(r"\n\s*\n+")


def iter_paragraphs(path, ext, content=None, include_comments=True):
    """Paragraph units for --context-paragraphs search.

    .docx uses the document's own paragraphs (plus each table row as one
    pseudo-paragraph and, by default, each comment); every other format splits
    its text on blank lines.
    """
    if ext == ".docx":
        from docx import Document

        doc = Document(path)
        paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                line = " | ".join(cell.text.strip() for cell in row.cells)
                if line.strip(" |"):
                    paras.append(line)
        if include_comments:
            paras.extend(docx_comments(doc))
        return paras
    if content is None:
        content = read_file_content(path, include_comments)
    if not content:
        return []
    return [p.strip() for p in _PARA_SPLIT.split(content.strip()) if p.strip()]


def _mark(text, q):
    """Wrap each case-insensitive occurrence of q in »…« (preserving case)."""
    low = text.lower()
    out = []
    i = 0
    ql = len(q)
    while i < len(text):
        j = low.find(q, i)
        if j == -1:
            out.append(text[i:])
            break
        out.append(text[i:j])
        out.append("»" + text[j : j + ql] + "«")
        i = j + ql
    return "".join(out)


def _match_header(label, count, noun="match"):
    plural = noun + ("es" if count != 1 else "")
    print(f"\n{'═' * 60}")
    print(f"  {label} ({count} {plural})")
    print(f"{'═' * 60}")


def search_lines(label, content, q, n):
    lines = content.splitlines()
    matching = [i for i, line in enumerate(lines) if q in line.lower()]
    if not matching:
        return 0
    _match_header(label, len(matching))
    shown = set()
    for idx in matching:
        start = max(0, idx - n)
        end = min(len(lines), idx + n + 1)
        for i in range(start, end):
            if i not in shown:
                marker = ">>>" if i == idx else "   "
                print(f"  {marker} {i + 1:4d} | {lines[i]}")
                shown.add(i)
        if end < len(lines):
            print("       ...")
    return len(matching)


def search_paragraphs(label, paras, q, n):
    matching = [i for i, p in enumerate(paras) if q in p.lower()]
    if not matching:
        return 0
    _match_header(label, len(matching))
    shown = set()
    for idx in matching:
        start = max(0, idx - n)
        end = min(len(paras), idx + n + 1)
        for i in range(start, end):
            if i not in shown:
                marker = ">>>" if i == idx else "   "
                body = _mark(paras[i], q) if i == idx else paras[i]
                print(f"  {marker} ¶{i + 1} | {body}")
                shown.add(i)
        print()
    return len(matching)


def search_chars(label, content, q, n):
    low = content.lower()
    ql = len(q)
    offsets = []
    start = 0
    while True:
        idx = low.find(q, start)
        if idx == -1:
            break
        offsets.append(idx)
        start = idx + ql
    if not offsets:
        return 0
    windows = []
    for off in offsets:
        s = max(0, off - n)
        e = min(len(content), off + ql + n)
        if windows and s <= windows[-1][1]:
            windows[-1] = (windows[-1][0], max(windows[-1][1], e))
        else:
            windows.append((s, e))
    _match_header(label, len(offsets))
    for s, e in windows:
        prefix = "…" if s > 0 else ""
        suffix = "…" if e < len(content) else ""
        snippet = _mark(content[s:e], q).replace("\n", " ")
        print(f"  {prefix}{snippet}{suffix}")
        print()
    return len(offsets)


def _render_rows(str_rows, wanted, match_rows):
    if not wanted:
        return ""
    ncols = max((len(str_rows[i]) for i in wanted), default=0)
    col_widths = []
    for c in range(ncols):
        w = 0
        for i in wanted:
            r = str_rows[i]
            if c < len(r):
                w = max(w, len(r[c]))
        col_widths.append(min(w, 40))
    out = []
    for i in sorted(wanted):
        r = str_rows[i]
        padded = [
            (r[c] if c < len(r) else "").ljust(col_widths[c]) for c in range(ncols)
        ]
        marker = ">>>" if i in match_rows else "   "
        out.append(f"  {marker} {i + 1:4d} | " + " | ".join(padded))
    return "\n".join(out)


def search_xlsx(label, path, query, n_rows, include_comments=True):
    """Spreadsheet-aware search: report Sheet!Coord + column/row headers per
    matching cell, then a window of the matching rows ± n_rows. Cell comments are
    searched too (unless include_comments is False)."""
    from openpyxl import load_workbook
    from openpyxl.utils import get_column_letter

    q = query.lower()
    wb = load_workbook(path, read_only=True, data_only=True)
    sections = []
    total = 0
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue
        header_row = rows[0]
        matches = []
        for r_idx, row in enumerate(rows):
            for c_idx, val in enumerate(row):
                if val is not None and q in str(val).lower():
                    matches.append((r_idx, c_idx, val))
        if not matches:
            continue
        lines = []
        for r_idx, c_idx, val in matches:
            coord = f"{get_column_letter(c_idx + 1)}{r_idx + 1}"
            bits = []
            if r_idx != 0 and c_idx < len(header_row) and header_row[c_idx] is not None:
                bits.append(f'col "{header_row[c_idx]}"')
            first = rows[r_idx][0] if rows[r_idx] else None
            if c_idx != 0 and first is not None:
                bits.append(f'row "{first}"')
            ctx = f"  [{' / '.join(bits)}]" if bits else ""
            lines.append(f"  >>> {sheet_name}!{coord}{ctx}  {val}")
        match_rows = {r for r, _, _ in matches}
        wanted = set()
        for r in match_rows:
            for i in range(max(0, r - n_rows), min(len(rows), r + n_rows + 1)):
                wanted.add(i)
        str_rows = [[str(c) if c is not None else "" for c in row] for row in rows]
        sections.append((sheet_name, lines, _render_rows(str_rows, wanted, match_rows)))
        total += len(matches)
    wb.close()
    comment_hits = [
        line
        for line in (xlsx_comments(path) if include_comments else [])
        if q in line.lower()
    ]
    total += len(comment_hits)
    if total == 0:
        return 0
    _match_header(label, total, noun="cell match")
    for sheet_name, lines, window in sections:
        print(f"\n  --- Sheet: {sheet_name} ---")
        for line in lines:
            print(line)
        print()
        print(window)
        print()
    if comment_hits:
        print("\n  --- Comments ---")
        for line in comment_hits:
            print(f"  >>> {_mark(line, q)}")
    return total


def search_file(label, path, ext, query, mode, n, n_rows, include_comments=True):
    """Dispatch a search over one file; print results, return match count."""
    if ext == ".xlsx":
        return search_xlsx(label, path, query, n_rows, include_comments)
    content = read_file_content(path, include_comments)
    if content is None:
        return 0
    q = query.lower()
    if mode == "paragraphs":
        paras = iter_paragraphs(path, ext, content, include_comments)
        return search_paragraphs(label, paras, q, n)
    if mode == "chars":
        return search_chars(label, content, q, n)
    return search_lines(label, content, q, n)


def resolve_context(args):
    """Return (mode, n) from the mutually exclusive context flags (default lines/2)."""
    for mode, value in (
        ("lines", args.context),
        ("paragraphs", args.context_paragraphs),
        ("chars", args.context_chars),
    ):
        if value is not None:
            return mode, value
    return "lines", 2
